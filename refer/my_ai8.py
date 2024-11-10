import os
import requests
import yfinance as yf
from datetime import datetime
from docx import Document

# 股票代碼清單
stock_list = ['7769.TW', '6988.TW', '3616.TW', '7760.TW']  # 注意加上 .TW 表示台灣股票
youtube_api_key = 'AIzaSyDirDw_Tw2lSp3rano1SLuXBqLgoR2J-zo'  # 替換為您的 YouTube Data API 金鑰

# 定義查詢影片的功能
def fetch_videos(stock_id):
    query = f"{stock_id} 台灣 股票"
    url = f"https://www.googleapis.com/youtube/v3/search?part=snippet&q={query}&key={youtube_api_key}&maxResults=5&regionCode=TW&type=video"
    response = requests.get(url)
    video_data = []
    
    if response.status_code == 200:
        video_items = response.json().get('items', [])
        for item in video_items:
            video_data.append({
                "title": item["snippet"]["title"],
                "date": item["snippet"]["publishedAt"],
                "link": f"https://www.youtube.com/watch?v={item['id']['videoId']}"
            })
    return video_data

# 獲取股票資訊
def fetch_stock_info(stock_id):
    stock = yf.Ticker(stock_id)
    stock_info = stock.info
    return {
        "symbol": stock_info.get("symbol"),
        "name": stock_info.get("longName"),
        "marketCap": stock_info.get("marketCap"),
        "previousClose": stock_info.get("regularMarketPreviousClose"),
        "currentPrice": stock_info.get("currentPrice"),
        "description": stock_info.get("longBusinessSummary")
    }

# 統整每支股票的報告內容
def compile_report(stock_id):
    videos = fetch_videos(stock_id)
    stock_info = fetch_stock_info(stock_id)
    
    goodinfo_link = f"https://goodinfo.tw/tw/StockDetail.asp?STOCK_ID={stock_id.split('.')[0]}&INIT=T"
    
    report_content = f"## 股票代碼: {stock_info['symbol']} ({stock_info['name']})\n"
    report_content += f"### Goodinfo連結:\n"
    report_content += f'<a href="{goodinfo_link}">{goodinfo_link}</a>\n\n'
    
    report_content += "### 股票資訊:\n"
    report_content += f"- 市值: {stock_info['marketCap']}\n"
    report_content += f"- 前收盤價: {stock_info['previousClose']}\n"
    report_content += f"- 當前價格: {stock_info['currentPrice']}\n"
    report_content += f"- 描述: {stock_info['description']}\n\n"

    report_content += "### 影片分析:\n"
    
    if videos:
        for v in videos:
            report_content += f"- **{v['date']}**: {v['title']}\n  [來源]({v['link']})\n\n"
    else:
        report_content += "- 無相關影片資料\n\n"

    return report_content[:1000]  # 限制字數在1000以內

# 生成 Word 報告
def generate_word(reports):
    os.makedirs("report_new", exist_ok=True)
    doc_name = datetime.now().strftime("%Y%m%d") + "_new.docx"
    doc_path = os.path.join("report_new", doc_name)
    
    doc = Document()
    
    for report in reports:
        # 將報告內容添加到 Word 文件中
        for line in report.split('\n'):
            if line.startswith("##"):
                doc.add_heading(line[3:], level=2)  # 添加二級標題
            elif line.startswith("###"):
                doc.add_heading(line[4:], level=3)  # 添加三級標題
            elif line.startswith("-"):
                doc.add_paragraph(line[2:], style='ListBullet')  # 添加項目符號列表
            else:
                doc.add_paragraph(line)  # 添加普通段落
        
        doc.add_page_break()  # 每個報告之間插入分頁
    
    doc.save(doc_path)
    print(f"報告已生成: {doc_path}")

# 主程式執行
if __name__ == "__main__":
    reports = []
    
    for stock_id in stock_list:
        report_content = compile_report(stock_id)
        reports.append(report_content)
    
    generate_word(reports)