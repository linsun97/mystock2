import os
import requests
# import mysql.connector
from datetime import datetime
from docx import Document
# from duckduckgo_search import ddg
from duckduckgo_search import DDGS
import mariadb

# 股票代碼清單
stock_ids = ['7769', '3616', '7760']  # 股票代碼，不包括 .TW 後綴
youtube_api_key = 'AIzaSyDirDw_Tw2lSp3rano1SLuXBqLgoR2J-zo'  # 替換為您的 YouTube Data API 金鑰

# 連接到 MariaDB 資料庫
def get_stock_names(stock_ids):
    connection = mariadb.connect(
        user="root",       # MariaDB 使用者名稱
        password="nineseve9173",   # MariaDB 密碼
        host="localhost",           # 資料庫主機位址，通常是 localhost
        port=3306,                  # 預設 MariaDB 的連接埠
        database="stock"            # 資料庫名稱
    )
    cursor = connection.cursor()
    
    stock_names = {}
    query = "SELECT stockid, stockname FROM all_id_name_shin WHERE stockid IN (%s)" % ','.join(['%s'] * len(stock_ids))
    cursor.execute(query, stock_ids)
    
    for (stockid, stockname) in cursor.fetchall():
        stock_names[stockid] = stockname
    
    cursor.close()
    connection.close()
    
    return stock_names

# 定義查詢新聞的功能
def fetch_news(stock_id, stock_name):
    query = f"{stock_id} {stock_name} 台股 掛牌 登錄 興櫃"
    print(query)
    results = DDGS().text(query, max_results=5)
    news_data = []
    
    for result in results:
        news_data.append({
            "title": result["title"],
            "date": result.get("date", "未知"),
            "summary": result["body"],
            "link": result["href"]
        })
    
    return news_data

# 定義查詢影片的功能
def fetch_videos(stock_id, stock_name):
    # query = f"{stock_id} {stock_name}"
    query = f"{stock_name}"
    print(query)
    url = f"https://www.googleapis.com/youtube/v3/search?part=snippet&q={query}&key={youtube_api_key}&maxResults=5&regionCode=TW&type=video"
    response = requests.get(url)
    video_data = []
    
    if response.status_code == 200:
        video_items = response.json().get('items', [])
        for item in video_items:
            video_data.append({
                "title": item["snippet"]["title"],
                "date": item["snippet"]["publishedAt"],
                "link": f"https://www.youtube.com/watch?v={item['id']['videoId']}"
            })
    return video_data

# 統整每支股票的報告內容
def compile_report(stock_id, stock_name):
    news = fetch_news(stock_id, stock_name)
    videos = fetch_videos(stock_id, stock_name)
    
    goodinfo_link = f"https://goodinfo.tw/tw/StockDetail.asp?STOCK_ID={stock_id}&INIT=T"
    
    report_content = f"## 股票代碼: {stock_id} ({stock_name})\n"
    report_content += f"### Goodinfo連結:\n"
    report_content += f'{goodinfo_link}\n\n'
    
    report_content += "### 新聞分析:\n"
    
    if news:
        for n in news:
            report_content += f"- **{n['date']}**: {n['title']}\n  {n['summary']}\n  [來源]({n['link']})\n\n"
    else:
        report_content += "- 無相關新聞資料\n\n"

    report_content += "### 影片分析:\n"
    
    if videos:
        for v in videos:
            report_content += f"- **{v['date']}**: {v['title']}\n  [來源]({v['link']})\n\n"
    else:
        report_content += "- 無相關影片資料\n\n"

    return report_content[:3000]  # 限制字數在3000以內

# 生成 Word 報告
def generate_word(reports):
    os.makedirs("report_new", exist_ok=True)
    doc_name = datetime.now().strftime("%Y%m%d") + "_new.docx"
    doc_path = os.path.join("report_new", doc_name)
    
    doc = Document()
    
    for report in reports:
        # 將報告內容添加到 Word 文件中
        for line in report.split('\n'):
            if line.startswith("##"):
                doc.add_heading(line[3:], level=2)  # 添加二級標題
            elif line.startswith("###"):
                doc.add_heading(line[4:], level=3)  # 添加三級標題
            elif line.startswith("-"):
                doc.add_paragraph(line[2:], style='ListBullet')  # 添加項目符號列表
            else:
                doc.add_paragraph(line)  # 添加普通段落
        
        doc.add_page_break()  # 每個報告之間插入分頁
    
    doc.save(doc_path)
    print(f"報告已生成: {doc_path}")

# 主程式執行
if __name__ == "__main__":
    stock_names_dict = get_stock_names(stock_ids)
    print(stock_names_dict)
    # quit()
    
    reports = []
    
    for  stock_id, stock_name in stock_names_dict.items():
        # print(stock_id)
        # stock_name = stock_names_dict.get(stock_id, "未知股票")
        # print(stock_name)
        # quit()
        report_content = compile_report(stock_id, stock_name)
        reports.append(report_content)
    
    generate_word(reports)