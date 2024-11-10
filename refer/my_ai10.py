import os
import requests
from datetime import datetime, timedelta
from docx import Document
# from duckduckgo_search import ddg
from duckduckgo_search import DDGS
import mariadb

# 給定的股票代碼 list
stock_ids = [7769, 3616, 7760]

# 建立一個空的 list 儲存股票名稱
stock_names = []

try:
    # 連接 MariaDB 資料庫
    connection = mariadb.connect(
        user="root",       # MariaDB 使用者名稱
        password="nineseve9173",   # MariaDB 密碼
        host="localhost",           # 資料庫主機位址，通常是 localhost
        port=3306,                  # 預設 MariaDB 的連接埠
        database="stock"            # 資料庫名稱
    )
    
    # 建立 cursor 物件
    cursor = connection.cursor()

    # 準備 SQL 查詢語句
    query = "SELECT stockname FROM all_id_name_shin WHERE stockid = ?"

    # 依序查詢每個股票代碼
    for stock_id in stock_ids:
        cursor.execute(query, (stock_id,))
        result = cursor.fetchone()
        if result:
            stock_names.append(result[0])  # 將查詢到的股票名稱加入 list
        else:
            stock_names.append("查無此股票")  # 若查無資料，加入占位訊息

except mariadb.Error as e:
    print(f"Error connecting to MariaDB Platform: {e}")
finally:
    # 關閉資料庫連線
    if connection:
        connection.close()

# 印出查詢結果
print(stock_names)
# quit()
# 股票代碼清單
stock_list = [f"{stock_id}.TW" for stock_id in stock_ids]



# stock_list = ['7769.TW', '6988.TW', '3616.TW', '7760.TW']  # 注意加上 .TW 表示台灣股票
youtube_api_key = 'AIzaSyDirDw_Tw2lSp3rano1SLuXBqLgoR2J-zo'  # 替換為您的 YouTube Data API 金鑰

# 定義查詢新聞的功能
def fetch_news(stock_id,stock_name):
    query = f"{stock_id} {stock_name} 台股 掛牌 登錄 興櫃"
    # results = ddg(query, max_results=5)
    results = DDGS().text(query, max_results=5)
    # results = DDGS().text("台灣股市", max_results=5)
    news_data = []
    
    for result in results:
        news_data.append({
            "title": result["title"],
            "date": result.get("date", "未知"),
            "summary": result["body"],
            "link": result["href"]
        })
    
    return news_data

# 定義查詢影片的功能
def fetch_videos(stock_id,stock_name):
    query = f"{stock_id} {stock_name} 台股 掛牌 登錄 前業績發表會"
    url = f"https://www.googleapis.com/youtube/v3/search?part=snippet&q={query}&key={youtube_api_key}&maxResults=5&regionCode=TW&type=video"
    response = requests.get(url)
    video_data = []
    
    if response.status_code == 200:
        video_items = response.json().get('items', [])
        for item in video_items:
            video_data.append({
                "title": item["snippet"]["title"],
                "date": item["snippet"]["publishedAt"],
                "link": f"https://www.youtube.com/watch?v={item['id']['videoId']}"
            })
    return video_data

# 統整每支股票的報告內容
def compile_report(stock_id,stock_name):
    news = fetch_news(stock_id,stock_name)
    videos = fetch_videos(stock_id,stock_name)
    
    goodinfo_link = f"https://goodinfo.tw/tw/StockDetail.asp?STOCK_ID={stock_id.split('.')[0]}&INIT=T"
    
    report_content = f"## 股票代碼: {stock_id}\n"
    report_content += f"### Goodinfo連結:\n"
    report_content += f'{goodinfo_link}\n\n'
    
    report_content += "### 新聞分析:\n"
    
    if news:
        for n in news:
            report_content += f"- **{n['date']}**: {n['title']}\n  {n['summary']}\n  [來源]({n['link']})\n\n"
    else:
        report_content += "- 無相關新聞資料\n\n"

    report_content += "### 影片分析:\n"
    
    if videos:
        for v in videos:
            report_content += f"- **{v['date']}**: {v['title']}\n  [來源]({v['link']})\n\n"
    else:
        report_content += "- 無相關影片資料\n\n"

    return report_content[:2000]  # 限制字數在1000以內

# 生成 Word 報告
def generate_word(reports):
    os.makedirs("report_new", exist_ok=True)
    doc_name = datetime.now().strftime("%Y%m%d") + "_new.docx"
    doc_path = os.path.join("report_new", doc_name)
    
    doc = Document()
    
    for report in reports:
        # 將報告內容添加到 Word 文件中
        for line in report.split('\n'):
            if line.startswith("##"):
                doc.add_heading(line[3:], level=2)  # 添加二級標題
            elif line.startswith("###"):
                doc.add_heading(line[4:], level=3)  # 添加三級標題
            elif line.startswith("-"):
                doc.add_paragraph(line[2:], style='ListBullet')  # 添加項目符號列表
            else:
                doc.add_paragraph(line)  # 添加普通段落
        
        doc.add_page_break()  # 每個報告之間插入分頁
    
    doc.save(doc_path)
    print(f"報告已生成: {doc_path}")

# 主程式執行
if __name__ == "__main__":
    reports = []
    
    for stock_id in stock_list:
        report_content = compile_report(stock_id)
        reports.append(report_content)
    
    generate_word(reports)